import streamlit as st
from config import Config, load_css
from file_processor import FileProcessor
from ai_analyzer import AIAnalyzer
from docx import Document
from io import BytesIO

# --- Setup ---
st.set_page_config(page_title=Config.PAGE_TITLE, page_icon=Config.PAGE_ICON, layout=Config.LAYOUT)
load_css()

# --- Helper Functions ---
def create_docx(content, topic):
    """Converts Markdown/Text content to a Word Document."""
    doc = Document()
    doc.add_heading(f'Study Guide: {topic}', 0)
    
    # Simple cleanup to make markdown look okay in Word
    lines = content.split('\n')
    for line in lines:
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            runner = p.add_run(line.replace('**', ''))
            runner.bold = True
        else:
            doc.add_paragraph(line)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Sidebar ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/2666/2666505.png", width=50)
    st.title("EconWiz Pro")
    
    # API KEY LOGIC
    api_key = None
    if "OPENAI_API_KEY" in st.secrets:
        api_key = st.secrets["OPENAI_API_KEY"]
        st.success(f"Connected: ...{api_key[-4:]}")
    else:
        st.error("Missing API Key")
        api_key = st.text_input("Enter OpenAI Key:", type="password")
    
    st.divider()
    
    st.header("🎓 Academic Level")
    level = st.select_slider("Select Difficulty", options=["High School", "Undergrad (Intro)", "Undergrad (Adv)", "Grad/PhD"])

# --- Main Interface ---
st.title(f"{Config.PAGE_ICON} {Config.PAGE_TITLE}")

# Tabs for different modes
tab1, tab2 = st.tabs(["🧮 Homework Solver", "📚 Study Guide Generator"])

# --- TAB 1: HOMEWORK SOLVER ---
with tab1:
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.info("Upload a screenshot or simply type your variables below.")
        
        # Topic Selector for Solver
        topic_solver = st.selectbox("Topic", [
            "Supply & Demand", "Consumer Theory", "Market Structures", 
            "GDP & Inflation", "Fiscal/Monetary Policy", "Game Theory", "Econometrics"
        ], key="solver_topic")
        
        # File Upload
        uploaded_file = st.file_uploader("Upload Problem (Image/PDF)", type=Config.ALLOWED_EXTENSIONS)
        
        # Text Input
        user_query = st.text_area(
            "Problem / Variables", 
            height=150, 
            placeholder="Paste your question OR enter variables:\nExample:\nQd = 100 - 2P\nQs = 10 + 4P\nFind Equilibrium Price."
        )
        
        solve_btn = st.button("🚀 Solve Problem", type="primary")

    with col2:
        if solve_btn:
            if not api_key:
                st.error("Please provide an API Key.")
            elif not (user_query or uploaded_file):
                st.warning("Please provide a file or text input.")
            else:
                with st.spinner("Crunching numbers..."):
                    # Process File
                    text_from_file = ""
                    img_parts = []
                    preview_img = None
                    
                    if uploaded_file:
                        processor = FileProcessor()
                        text_from_file, img_parts, preview_img = processor.process_file(uploaded_file)
                        if preview_img:
                            st.image(preview_img, caption="Problem Image", width=300)
                    
                    # Combine
                    full_query = user_query
                    if text_from_file:
                        full_query += f"\n\n[FILE DATA]:\n{text_from_file[:1500]}"
                    
                    # Analyze
                    analyzer = AIAnalyzer(api_key, Config.MODEL_NAME)
                    result = analyzer.analyze_problem(img_parts, full_query, topic_solver, level)
                    
                    # Display
                    st.markdown("### Solution")
                    st.markdown(result)

# --- TAB 2: STUDY GUIDE GENERATOR ---
with tab2:
    st.header("Create Custom Study Material")
    st.write("Generate a detailed study guide for any concept and export it to Word.")
    
    sg_col1, sg_col2 = st.columns([1, 2])
    
    with sg_col1:
        custom_topic = st.text_input("Enter Topic for Study Guide", placeholder="e.g. Nash Equilibrium, IS-LM Model")
        generate_btn = st.button("✨ Generate Guide", use_container_width=True)
    
    with sg_col2:
        if generate_btn and custom_topic and api_key:
            with st.spinner(f"Drafting study guide for '{custom_topic}'..."):
                analyzer = AIAnalyzer(api_key, Config.MODEL_NAME)
                guide_content = analyzer.generate_study_guide(custom_topic, level)
                
                # Store in session state so it persists
                st.session_state['guide_content'] = guide_content
                st.session_state['guide_topic'] = custom_topic

        # Display Result if available
        if 'guide_content' in st.session_state:
            st.markdown("---")
            st.subheader(f"📖 {st.session_state['guide_topic']}")
            st.markdown(st.session_state['guide_content'])
            
            # Export Button
            docx_file = create_docx(st.session_state['guide_content'], st.session_state['guide_topic'])
            
            st.download_button(
                label="📥 Download as Word Doc (.docx)",
                data=docx_file,
                file_name=f"Study_Guide_{st.session_state['guide_topic'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

# --- Footer ---
st.divider()
st.caption("EconWiz Pro | Powered by OpenAI GPT-4o")
